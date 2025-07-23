import os
from docx.shared import Inches
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def get_images_from_dir(path="../assets"):
    all_files = os.listdir(path)
    all_images = [f for f in all_files if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
    return [os.path.join(path, f) for f in all_images]

def create_table_photos(document, last_position, list_of_images_path):
    title = document.add_paragraph("Registros Fotográficos das Não Conformidades")
    title.style = 'Arial10'
    photos_table = document.add_table(rows=6, cols=2)
    for i in range(len(list_of_images_path)):
        photo_line = i // 2 * 2
        subtitle_line = photo_line + 1
        collum = i % 2

        photos_table.cell(photo_line, collum).paragraphs[0].add_run().add_picture(list_of_images_path[i], width=Inches(2.2))
        photos_table.cell(subtitle_line, collum).text = "Legenda da foto"

    last_position._element.addnext(title._element)
    title._element.addnext(photos_table._element)
    set_borders_photos(photos_table)
    print(">>> Tabela de Fotos Criada")

    return photos_table

def divide_photos(document, last_position, list_of_images_path):
    for i in range(0, len(list_of_images_path), 6):
        table_photos = list_of_images_path[i:i+6]
        last_position = create_table_photos(document, last_position, table_photos)
        print(f">>> {i}° bloco concluido")

def set_borders_photos(table):
    tbl = table._element
    tblPr = tbl.xpath('./w:tblPr')[0]

    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')   
        border.set(qn('w:sz'), '8')        
        border.set(qn('w:space'), '0')       
        border.set(qn('w:color'), '000000')   
        tblBorders.append(border)

    tblPr.append(tblBorders)