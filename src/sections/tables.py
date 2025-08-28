from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from unidecode import unidecode
from excel import get_non_conformities, get_inspections_data, list_of_all_units, documents_excel, town_statistics
from utils import search_paragraph, apply_background_color, set_column_widths
from images import set_borders_table  
import pandas as pd

def create_generic_nx2_table(document, rows_data, text_after_paragraph, col1_width, col2_width, align_left=False):
    """
    Cria uma tabela nx2 a partir de uma lista de tuplas (key, value).
    
    rows_data: lista de tuplas (key, value), value="" indica subtópico
    text_after_paragraph: Insere a tabela após o parágrafo com esse texto
    align_left: se True, sobrescreve o alinhamento das células para LEFT
    """
    table = document.add_table(rows=0, cols=2)
    
    for key, value in rows_data:
        row_cells = table.add_row().cells
        if value == "":
            cell = row_cells[0].merge(row_cells[1])
            format_header_cell(cell, key)
        else:
            format_data_cell(row_cells[0], key)
            for run in row_cells[0].paragraphs[0].runs:
                run.bold = True
            format_data_cell(row_cells[1], value)
    
    if align_left:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    set_borders_table(table)
    set_column_widths(table, col1_width, col2_width)
    paragraph_index = search_paragraph(document, text_after_paragraph)[0]
    document.paragraphs[paragraph_index]._element.addnext(table._element)

def create_general_information_table(document, text):
    report_data = get_inspections_data()
    general_info = [
        ("3.1 DO TITULAR", ""),
        ("Titular:", "Microrregião de Água e Esgoto RMR-PAJEÚ/Microrregião de Água e Esgoto SERTÃO"),
        ("Endereço:", "Avenida Cruz Cabugá, 1387 - Santo Amaro - Recife, PE - CEP: 50040-905"),
        ("Responsável:", "Artur Paiva Coutinho"),
        ("Município:", report_data["Municipio"]),
        ("3.2 DO REGULADO", ""),
        ("Regulado:", "Companhia Pernambucana de Saneamento - Compesa"),
        ("Responsável:", "Dr. Alex Machado Campos"),
        ("Endereço:", "Av. Cruz Cabugá, 1387 - Santo Amaro - Recife, PE - CEP: 50040-905"),
        ("Representantes por acompanhar:", report_data["Representantes por acompanhar"]),
        ("3.3 DO REGULADOR", ""),
        ("Regulador:", "Agência de Regulação de Pernambuco"),
        ("Diretor Presidente:", "Carlos Porto Filho"),
        ("Endereço:", "Avenida Conselheiro Rosa e Silva, 975, Aflitos, Recife/PE, CEP: 52.050-020."),
        ("Responsáveis pela fiscalização:", f"{report_data['Analista 1']} e {report_data['Analista 2']}"),
        ("Período da Fiscalização:", report_data["Periodo"]),
        ("Tipo de Fiscalização:", "Direta e periódica.")
    ]

    create_generic_nx2_table(document, general_info, text, col1_width=1, col2_width=9, align_left=True)

def create_abbreviations_table(document, text):
    abbreviations = [
        ("Sigla", "Definição"),
        ("ETA", "Estação de Tratamento de Água"),
        ("ETE", "Estação de Tratamento de Esgoto"),
        ("EEab", "Estação Elevatória de água bruta"),
        ("EEat", "Estação Elevatória de água tratada"),
        ("REL", "Reservatório Elevado"),
        ("RAP", "Reservatório Apoiado"),
        ("CMB", "Conjunto Moto Bomba"),
        ("GNR", "Gerência de Unidade de Negócios Regional"),
        ("SAA", "Sistemas de Abastecimento de Água"),
        ("SES", "Sistemas de Esgotamento Sanitário"),
        ("IUA", "Índice de Universalização do Abastecimento de Água"),
        ("IUE",	"Índice de Universalização de Coleta de Esgotos Sanitários"),
        ("IUT", "Índice de Universalização de Tratamento de Esgotos Sanitários"),
        ("ICA",	"Índice de Cobertura de Abastecimento de Água"),
        ("ICE", "Índice de Cobertura de Esgotamento Sanitário"),
        ("IPD","Índice de Perdas na Distribuição"),
        ("IQAP","Índice da Qualidade da Água Potável")
    ]

    create_generic_nx2_table(document, abbreviations, text, col1_width=2, col2_width=8, align_left=True)
    
def create_documents_table(document, text):
    df_documents = documents_excel.copy()
    table = document.add_table(rows=1, cols=len(df_documents.columns))
    set_column_widths(table, 6.5, 0.5, 0.5, 6.5)

    for idx, col_name in enumerate(df_documents.columns):
        format_header_cell(table.rows[0].cells[idx], col_name, font_size=10)
 
    for _, row in df_documents.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            format_data_cell(cells[idx], value, font_size=10)

    set_borders_table(table)
    paragraph_index = search_paragraph(document, text)[0]
    document.paragraphs[paragraph_index]._element.addnext(table._element)

    print("✅ Tabela de documentos criada.")

def create_town_units_table(document, text):
    report_data = get_inspections_data()

    town_name = unidecode(str(report_data["Municipio"]).strip().lower())
    inspection_type = unidecode(str(report_data["Tipo da Fiscalização"]).strip().lower())

    units_df = list_of_all_units.copy()
    units_df.columns = units_df.columns.str.strip()

    units_df["MUNICÍPIO_norm"] = units_df["MUNICÍPIO"].astype(str).str.strip().str.lower().apply(unidecode)
    units_df["AGUA_ESGOTO_norm"] = units_df["ÁGUA/ESGOTO"].astype(str).str.strip().str.lower().apply(unidecode)

    filtered_units = units_df[units_df["MUNICÍPIO_norm"] == town_name]

    if "esgoto" in inspection_type:
        filtered_units = filtered_units[filtered_units["AGUA_ESGOTO_norm"].str.contains("esgoto")]
    else:
        filtered_units = filtered_units[filtered_units["AGUA_ESGOTO_norm"].str.contains("agua")]

    filtered_units.insert(0, "ITEM", range(1, len(filtered_units) + 1))
    filtered_units["OBSERVAÇÃO"] = ""

    final_df = filtered_units[["ITEM", "SISTEMA", "UNIDADE", "OBSERVAÇÃO"]]
    table = document.add_table(rows=1, cols=len(final_df.columns))
    set_column_widths(table, 0.3, 4.0, 6.0, 1.7)

    for idx, col_name in enumerate(final_df.columns):
        format_header_cell(table.rows[0].cells[idx], col_name, font_size=10)

    for _, row in final_df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            format_data_cell(cells[idx], value, font_size=10)

    set_borders_table(table)
    paragraph_index = search_paragraph(document, text)[0]
    document.paragraphs[paragraph_index]._element.addnext(table._element)

    print("✅ Tabela de unidades do município criada.")    

def create_statistics_table(document, text):
    df_statistics = town_statistics.copy()
    report_data = get_inspections_data()
    report_town = report_data["Municipio"].upper()

    pernambuco_stats = {
        "Quantidade de economias residenciais ativas de água (A) - EAA": "2.261.695",
        "Quantidade de economias residenciais inativas de água (B)-EIA": "377.745",
        "Quantidade de economias residenciais ativas de esgoto (C) - EAE": "654.143",
        "Quantidade de economias residenciais inativas de esgoto (D) - EIE": "300.683",
        "Quantidade de economias residenciais ativas com tratamento de esgoto (E) - EAT": "654.143",
        "Quantidade de economias residenciais inativas com tratamento de esgoto (F) - EIT": "300.683",
        "Quantidade de domicílios residenciais existentes na área de abrangência do prestador de serviços (G) - DAP": "2.646.895"
    }

    selected_columns = list(pernambuco_stats.keys())

    stats_from_this_town = df_statistics[df_statistics["ANO BASE: 2023"] == report_town]
    stats_from_this_town = stats_from_this_town[selected_columns]

    table = document.add_table(rows=len(selected_columns) + 1, cols=3)
    set_column_widths(table, 4, 1.5, 1.5)

    format_header_cell(table.rows[0].cells[0], "INFORMAÇÃO", font_size=10)
    format_header_cell(table.rows[0].cells[1], "PERNAMBUCO", font_size=10)
    format_header_cell(table.rows[0].cells[2], report_town, font_size=10)

    for idx, indicador in enumerate(selected_columns, start=1):

        format_data_cell(table.rows[idx].cells[0], indicador, font_size=10)

        format_data_cell(table.rows[idx].cells[1], pernambuco_stats[indicador], font_size=10)

        valor_municipio = stats_from_this_town.iloc[0][indicador] if not stats_from_this_town.empty else ""
        if isinstance(valor_municipio, (int, float)):
            valor_municipio = str(int(valor_municipio))

        format_data_cell(table.rows[idx].cells[2], valor_municipio, font_size=10)

    set_borders_table(table)
    paragraph_index = search_paragraph(document, text)[0]
    document.paragraphs[paragraph_index]._element.addnext(table._element)

    print("✅ Tabela de estatísticas criada.")

def create_quality_index_table(document: Document, text: str):
    df_statistics = town_statistics.copy()
    report_data = get_inspections_data()
    report_town = report_data["Municipio"].upper()

    df_statistics.columns = df_statistics.columns.str.replace(r"\s+\(\%\)", "(%)", regex=True)
    df_statistics.columns = df_statistics.columns.str.strip()

    selected_columns = ["IUA(%)", "IUE(%)", "IUT(%)", "ICA", "ICE", "IPD", "IQAP"]

    stats_from_this_town = df_statistics[df_statistics["ANO BASE: 2023"] == report_town]

    valores_tabela = {}
    for col in selected_columns:
        if col in ["IUA(%)", "IUE(%)", "IUT(%)"]:
            valores_tabela[col] = (
                stats_from_this_town.iloc[0][col]
                if not stats_from_this_town.empty and col in stats_from_this_town.columns
                else ""
            )
        else:
            valores_tabela[col] = "-"

    table = document.add_table(rows=2, cols=len(selected_columns) + 1)

    format_header_cell(table.rows[0].cells[0], "Município", font_size=8)
    for col_idx, col_name in enumerate(selected_columns, start=1):
        format_header_cell(table.rows[0].cells[col_idx], col_name, font_size=8)

    format_data_cell(table.rows[1].cells[0], report_town, font_size=8)
    for col_idx, col_name in enumerate(selected_columns, start=1):
        format_data_cell(table.rows[1].cells[col_idx], valores_tabela[col_name], font_size=8)

    set_borders_table(table)

    paragraph_index = search_paragraph(document, text)[0]
    document.paragraphs[paragraph_index]._element.addnext(table._element)

    print(f"✅ Tabela de índices de qualidade criada para {report_town}")


def create_non_conformities_table(document: Document, text):
    df_ncs = get_non_conformities()

    selected_columns = [
        "Unidade", 
        "Não Conformidade", 
        "Nome da Foto", 
        "Artigo", 
        "Enquadramento", 
        "Determinações"
    ]
    df_ncs = df_ncs[selected_columns]

    table = document.add_table(rows=1, cols=len(selected_columns))
    set_column_widths(table, 3, 3, 0.5, 0.3, 3, 3)

    for idx, col_name in enumerate(selected_columns):
        format_header_cell(table.rows[0].cells[idx], col_name.upper(), font_size=8)

    for _, row in df_ncs.iterrows():
        cells = table.add_row().cells
        for idx, col_name in enumerate(selected_columns):
            format_data_cell(cells[idx], row[col_name], font_size=8)

    list_paragraph_index = search_paragraph(document, text)[0]
    document.paragraphs[list_paragraph_index]._element.addnext(table._element)

    set_borders_table(table)
    print("✅ Tabela de não conformidades criada.")


def format_header_cell(cell, text, font_size=10, font_name="Arial", bg_color="D9D9D9"):
    """Formata célula de cabeçalho: negrito, centralizado e fundo colorido."""
    para = cell.paragraphs[0]
    para.text = ""
    run = para.add_run(str(text))
    run.bold = True
    run.font.name = font_name
    run.font.size = Pt(font_size)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell._tc.get_or_add_tcPr().append(apply_background_color(bg_color))

def format_data_cell(cell, value, font_size=10, font_name="Arial"):
    """Formata célula de dados: centralizado e sem 'nan'."""
    para = cell.paragraphs[0]
    para.text = ""
    valor_texto = "" if pd.isna(value) else str(value)
    run = para.add_run(valor_texto)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER