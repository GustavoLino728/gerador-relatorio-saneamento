import os
from datetime import datetime, date
from excel import get_inspections_data
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

REPORT_DIR = "../reports/"
REPORT_NAME = "RELATÓRIO MODELO"
REPORT_EXT = ".docx"

def next_filename():
    for i in range(1, 1000):
        name = f"{REPORT_NAME}{'' if i == 0 else f' - {i}'}{REPORT_EXT}"
        path = os.path.join(REPORT_DIR, name)
        if not os.path.exists(path):
            return path

def search_paragraph(document, text):
    paragraphs_found_by_search = []
    print(">>> Busca no documento iniciada")
    for i, p in enumerate(document.paragraphs):
        if text in p.text:
            paragraphs_found_by_search.append(i)
            print(">>> Uma ocorrencia foi encontrada")
    return paragraphs_found_by_search

def substitute_placeholders(document):
    excel_data = get_inspections_data()

    def format_value(value):
        if value is None:
            return ""
        if isinstance(value, float) and value.is_integer():
            return str(int(value))
        if isinstance(value, (datetime, date)):
            return value.strftime("%d/%m/%Y")
        return str(value)

    def replace_in_runs(paragraph):
        new_runs = []
        for run in paragraph.runs:
            run_text = run.text
            replaced = False
            for key, value in excel_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in run_text:
                    run_text = run_text.replace(placeholder, format_value(value))
                    replaced = True
            if replaced:
                run.text = run_text
                run.font.color.rgb = None  


    for p in document.paragraphs:
        replace_in_runs(p)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_runs(p)

                
def apply_background_color(color_hex: str):
    """
    Cria um elemento de sombreado de célula com cor de fundo.
    Ex: 'D9D9D9' para cinza claro.
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    return shading_elm

def set_column_widths(table, *widths):
    """
    Ajusta a largura das colunas de uma tabela do python-docx.
    
    Parâmetros:
        table (docx.table.Table): Tabela a ser formatada.
        *widths (float): Larguras das colunas em polegadas. 
                         Passar um valor para cada coluna.
    """
    for row in table.rows:
        for col_idx, width in enumerate(widths):
            if col_idx < len(row.cells):
                row.cells[col_idx].width = Inches(width)
