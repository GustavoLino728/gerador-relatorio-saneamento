import os
from PIL import Image
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from excel import get_non_conformities 

def get_images_from_dir(path="../assets"):
    all_files = os.listdir(path)
    all_images = [f for f in all_files if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
    return [os.path.join(path, f) for f in all_images]

# def resize_images(path="../assets", size=(500, 500)):
#     for filename in os.listdir(path):
#         if filename.lower().endswith(('.jpg', '.jpeg', '.png')):
#             image_path = os.path.join(path, filename)
#             img = Image.open(image_path)
#             img_resized = img.resize(size, Image.LANCZOS)
#             img_resized.save(image_path)
#     print(">>> Imagens redimensionadas para 210x210")

def create_table_images(document, last_position, list_of_images_path, df_non_conformities):

    # Cria os elementos de parágrafo (vazios)
    space_before = document.add_paragraph()
    space_after = document.add_paragraph()
    space_before_element = space_before._element
    space_after_element = space_after._element

    title = document.add_paragraph("Registros Fotográficos das Não Conformidades")
    title.style = 'Arial10'
    images_table = document.add_table(rows=6, cols=2)
    
    for i in range(len(list_of_images_path)):
        image_line = i // 2 * 2
        subtitle_line = image_line + 1
        collum = i % 2

        images_table.cell(image_line, collum).paragraphs[0].add_run().add_picture(list_of_images_path[i], width=Inches(4.3))
        image_name = os.path.splitext(os.path.basename(list_of_images_path[i]))[0]
        line = df_non_conformities[df_non_conformities["Nome da Foto"] == image_name]

        if not line.empty:
            unit = line.iloc[0]["Unidade"]
            description = line.iloc[0]["Não Conformidade"]
            subtitle = f"{image_name} - {unit}: {description}"
        else:
            subtitle = f"{image_name} - NÃO ENCONTRADO"

        cell = images_table.cell(subtitle_line, collum)
        paragraph = cell.paragraphs[0]
        paragraph.style = 'Arial10'
        run = paragraph.add_run(subtitle)
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    last_position._element.addnext(space_before_element)
    space_before_element.addnext(title._element)
    title._element.addnext(space_after_element)
    space_after_element.addnext(images_table._element)

    set_borders_table(images_table)
    print(">>> Tabela de Fotos Criada")

    return images_table


def divide_images(document, last_position, list_of_images_path):
    for i in range(0, len(list_of_images_path), 6):
        table_images = list_of_images_path[i:i+6]
        last_position = create_table_images(document, last_position, table_images, get_non_conformities())
        print(f">>> {i}° bloco concluido")

def set_borders_table(table):
    """
    Adiciona bordas visíveis a uma tabela do python-docx.
    Pode ser usada para qualquer tabela de documento Word.
    """
    tbl = table._element

    # Verifica se <w:tblPr> existe, senão cria
    tblPr_list = tbl.xpath('./w:tblPr')
    if tblPr_list:
        tblPr = tblPr_list[0]
    else:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Cria as bordas
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')    # tipo da borda
        border.set(qn('w:sz'), '8')          # espessura
        border.set(qn('w:space'), '0')       # espaço entre borda e conteúdo
        border.set(qn('w:color'), '000000')  # cor preta
        tblBorders.append(border)

    tblPr.append(tblBorders)