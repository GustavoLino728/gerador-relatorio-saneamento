from docx import Document
from docx.shared import Inches
from photos import list_of_images_path

document = Document(r"../data/RELATÓRIO MODELO.docx")
position_paragraph = []

def create_table_photos(document, last_position, list_of_images_path):
    title = document.add_paragraph("Registros Fotográficos das Não Conformidades")
    photos_table = document.add_table(rows=6, cols=2)
    for i in range(len(list_of_images_path)):
        photo_line = i // 2 * 2
        subtitle_line = photo_line + 1
        collum = i % 2

        photos_table.cell(photo_line, collum).paragraphs[0].add_run().add_picture(list_of_images_path[i], width=Inches(2.2))
        photos_table.cell(subtitle_line, collum).text = "Legenda da foto"

    last_position._element.addnext(title._element)
    title._element.addnext(photos_table._element)

    return photos_table

for i, p in enumerate(document.paragraphs):
    if "APÊNDICE 1 - NÃO CONFORMIDADES" in p.text:
        position_paragraph.append(i)

last_position = document.paragraphs[position_paragraph[-1]]

for i in range(0, len(list_of_images_path), 6):
    table_photos = list_of_images_path[i:i+6]
    last_position = create_table_photos(document, last_position, table_photos)

document.save("..\\data\\RELATÓRIO MODELO.docx")    

