import pandas as pd
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from common.excel import get_non_conformities, get_inspections_data, units_df
from common.utils import search_paragraph, apply_background_color, set_column_widths, format_dict_values, set_table_margins, sanitize_value, set_borders_table, to_rows_data
from common.paths import SHEET_PATH


def create_generic_table(document, rows_data, text_after_paragraph, col_widths=None,
                         cell_padding=0.1, align_left=False, font_size=10):
    """
    Cria uma tabela genérica a partir de rows_data (lista de listas).

    rows_data:
        - Primeira linha: cabeçalho (colunas)
        - Demais linhas: dados
        - Se uma linha tiver menos valores que o número de colunas ou apenas 1 valor:
          será tratada como subtítulo (mescla todas as colunas)
          
    text_after_paragraph: posição para inserir a tabela
    col_widths: lista de larguras para cada coluna (ex: [2, 6, 3])
    cell_padding: preenchimento interno das células
    align_left: se True, alinha conteúdo à esquerda
    font_size: tamanho da fonte do conteúdo
    """
    if not rows_data:
        print("⚠️ Nenhum dado para criar tabela.")
        return

    n_cols = max(len(row) for row in rows_data)
    table = document.add_table(rows=0, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows_data):
        row_cells = table.add_row().cells

        if len(row) == 1 or len(row) < n_cols:
            merged_cell = row_cells[0]
            for c in row_cells[1:]:
                merged_cell = merged_cell.merge(c)
            format_header_cell(merged_cell, row[0], font_size=font_size)
            set_table_margins(merged_cell, top=cell_padding, bottom=cell_padding,
                              start=cell_padding, end=cell_padding)

        elif i == 0:
            for j, value in enumerate(row):
                format_header_cell(row_cells[j], value, font_size=font_size)
                set_table_margins(row_cells[j], top=cell_padding, bottom=cell_padding,
                                  start=cell_padding, end=cell_padding)

        else:
            for j, value in enumerate(row):
                format_data_cell(row_cells[j], value, font_size=font_size)
                set_table_margins(row_cells[j], top=cell_padding, bottom=cell_padding,
                                  start=cell_padding, end=cell_padding)

    if align_left:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    set_borders_table(table)
    if col_widths:
        set_column_widths(table, *col_widths)

    paragraph_index = search_paragraph(document, text_after_paragraph)[0]
    document.paragraphs[paragraph_index]._element.addnext(table._element)


def create_abbreviations_table(document, text):
    """
    Cria a tabela de siglas mostrando o significado de cada sigla que pode aparecer no relatorio (FIXO)
    Formato: 20 linhas x 2 colunas.
    Linha 1: Cabeçalho (Sigla, Definição)
    Linha 2-20: Siglas e Seus Significados
    """
    abbreviations = [
        ["Sigla", "Definição"],
        ["ETA", "Estação de Tratamento de Água"],
        ["ETE", "Estação de Tratamento de Esgoto"],
        ["EEab", "Estação Elevatória de água bruta"],
        ["EEat", "Estação Elevatória de água tratada"],
        ["REL", "Reservatório Elevado"],
        ["RAP", "Reservatório Apoiado"],
        ["CMB", "Conjunto Moto Bomba"],
        ["GNR", "Gerência de Unidade de Negócios Regional"],
        ["SAA", "Sistemas de Abastecimento de Água"],
        ["SES", "Sistemas de Esgotamento Sanitário"],
        ["IUA", "Índice de Universalização do Abastecimento de Água"],
        ["IUE", "Índice de Universalização de Coleta de Esgotos Sanitários"],
        ["IUT", "Índice de Universalização de Tratamento de Esgotos Sanitários"],
        ["ICA", "Índice de Cobertura de Abastecimento de Água"],
        ["ICE", "Índice de Cobertura de Esgotamento Sanitário"],
        ["IPD", "Índice de Perdas na Distribuição"],
        ["IQAP", "Índice da Qualidade da Água Potável"],
        ["NBR", "Normas Brasileiras"],
        ["CSAN", "Coordenadoria de Saneamento da ARPE"]
    ]
    
    create_generic_table(document, abbreviations, text, col_widths=[1, 9], align_left=True)
    
    
def create_general_information_table(document, text):
    """
    Cria a tabela de Informações gerais da fiscalização. Sobre o regulador, o regulado e o titular
    Formato: 17 linhas x 2 colunas.
    Linha 1/6/11: Titulos (3.1 DO TITULAR, 3.2 DO REGULADO, 3.3 DO REGULADOR)
    """
    report_data = get_inspections_data()
    
    general_info = [
        ["3.1 DO TITULAR"],
        ["Titular:", "Microrregião de Água e Esgoto RMR-PAJEÚ/Microrregião de Água e Esgoto SERTÃO"],
        ["Endereço:", "Avenida Cruz Cabugá, 1387 - Santo Amaro - Recife, PE - CEP: 50040-905"],
        ["Responsável:", "Artur Paiva Coutinho"],
        ["Município:", report_data["Municipio"]],
        ["3.2 DO REGULADO"],
        ["Regulado:", "Companhia Pernambucana de Saneamento - Compesa"],
        ["Responsável:", "Dr. Alex Machado Campos"],
        ["Endereço:", "Av. Cruz Cabugá, 1387 - Santo Amaro - Recife, PE - CEP: 50040-905"],
        ["Representantes por acompanhar:", report_data["Representantes por acompanhar"]],
        ["3.3 DO REGULADOR"],
        ["Regulador:", "Agência de Regulação de Pernambuco"],
        ["Diretor Presidente:", "Carlos Porto Filho"],
        ["Endereço:", "Avenida Conselheiro Rosa e Silva, 975, Aflitos, Recife/PE, CEP: 52.050-020."],
        ["Responsáveis pela fiscalização:", f"{report_data['Analista 1']} e {report_data['Analista 2']}"],
        ["Período da Fiscalização:", report_data["Período da Fiscalização"]],
        ["Tipo de Fiscalização:", "Direta e periódica."]
    ]
    
    create_generic_table(document, rows_data=general_info, text_after_paragraph=text, col_widths=[1, 9], align_left=True)


def create_documents_table(document, text):
    """
    Cria a tabela 1 - relativa as documentações necessarias para o processo, se elas foram enviadas ou não e porquê
    Formato: 12 linhas x 4 colunas.
    Linha 1: cabeçalho (DOCUMENTAÇÃO, SIM, NÃO, OBSERVAÇÕES)
    Linha 2-12: valores correspondentes
    """
    
    report_data = get_inspections_data()
    report_data["Tipo da Fiscalização"] = sanitize_value(report_data["Tipo da Fiscalização"])
    if report_data["Tipo da Fiscalização"] == "agua":
        documents_excel = pd.read_excel(SHEET_PATH, sheet_name='Envio de Documentos', header=1, nrows=11)
    if report_data["Tipo da Fiscalização"] == "esgoto":
        documents_excel = pd.read_excel(SHEET_PATH, sheet_name='Envio de Documentos', header=15)
    df_documents = documents_excel.copy()
    table = document.add_table(rows=1, cols=len(df_documents.columns))
    set_column_widths(table, 6.5, 0.5, 0.5, 6.5)

    for idx, col_name in enumerate(df_documents.columns):
        format_header_cell(table.rows[0].cells[idx], col_name, font_size=10)
 
    for _, row in df_documents.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            format_data_cell(cells[idx], value, font_size=10)

    set_borders_table(table)
    paragraph_index = search_paragraph(document, text)[0]
    document.paragraphs[paragraph_index]._element.addnext(table._element)


def create_town_units_table(document, text):
    """
    Cria a tabela 2 - Lista de Todas as Unidades do Município
    com base na planilha 'Cadastrar Unidades'.

    Estrutura esperada da planilha (linha 4 como cabeçalho):
    Municipio | Sistema | Tipo | Unidade | Observação
    """

    report_data = get_inspections_data()
    report_town = sanitize_value(report_data["Municipio"])
    inspection_type = sanitize_value(report_data["Tipo da Fiscalização"])

    units_df.columns = units_df.columns.str.strip()

    units_df["MUNICIPIO_NORMALIZED"] = units_df["Municipio"].apply(sanitize_value)
    units_df["TIPO_NORMALIZED"] = units_df["Tipo"].apply(sanitize_value)

    filtered_units = units_df[units_df["MUNICIPIO_NORMALIZED"] == report_town]

    if "agua" in inspection_type:
        allowed_types = ["eea", "eeab", "eeat", "eta", "rel/rap", "rel", "rap", "poço", "poco"]
    elif "esgoto" in inspection_type:
        allowed_types = ["eee", "ete"]
    else:
        print(f"⚠️ Tipo de fiscalização não reconhecido: {inspection_type}")
        return

    filtered_units = filtered_units[filtered_units["TIPO_NORMALIZED"].isin(allowed_types)]

    if filtered_units.empty:
        print("⚠️ Nenhuma unidade encontrada para este município/tipo de fiscalização.")
        return

    filtered_units = filtered_units.sort_values(by="Unidade", key=lambda col: col.str.lower())
    
    filtered_units.insert(0, "ITEM", range(1, len(filtered_units) + 1))
    final_df = filtered_units[["ITEM", "Sistema", "Unidade", "Observação"]]
    final_df = final_df.rename(columns={
        "Sistema": "SISTEMA",
        "Unidade": "UNIDADE",
        "Observação": "OBSERVAÇÃO"
    })

    rows_data = [final_df.columns.tolist()] 
    for row in final_df.itertuples(index=False, name=None):
        rows_data.append(list(row))
        
    create_generic_table(document=document, rows_data=rows_data, text_after_paragraph=text, col_widths=[0.8, 4, 6, 2], align_left=False)


def create_last_report_table(document, text):
    """
    Cria a tabela 3 - Com informações relevantes a última fiscalização naquele municipio
    Formato: 5 linhas x 2 colunas.
    Linha 1: cabeçalho (Contexto)
    Linha 2: Ultima Fiscalização
    Linha 3: Total NCs da Ultima Fiscalização
    Linha 4: DESDOBRAMENTOS
    Linha 5: NCs RESIDUAIS
    """
    report_data = format_dict_values(get_inspections_data())
    last_report_data = {
        "ÚLTIMA FISCALIZAÇÃO": report_data["Ultima Fiscalização (Data)"],
        "TOTAL DE NCs DA ÚLTIMA FISCALIZAÇÂO": report_data["Total NCS UF"],
        "DESDOBRAMENTOS": report_data["Desdobramentos"],
        "NCs RESIDUAIS": report_data["NCS Residuais"]
    }
    
    rows = to_rows_data(last_report_data, subtitle="CONTEXTO")
    create_generic_table(document, rows, text, col_widths=[2.5, 5], cell_padding=0.4, align_left=True)


def create_non_conformities_table(document, text):
    """
    Cria a tabela de não conformidades.
    Formato: 1+N linhas x 6 colunas.
    Linha 1: cabeçalho (Unidade, Não Conformidad, Nome da Foto, Artigo, Enquadramento, Determinações)
    Linha 2-N: Dados correspondentes
    """
    df_ncs = get_non_conformities()

    selected_columns = [
        "Unidade", 
        "Não Conformidade", 
        "Nome da Foto", 
        "Artigo", 
        "Enquadramento", 
        "Determinações"
    ]

    df_ncs = df_ncs[selected_columns].fillna("")
    rows_data = to_rows_data([df_ncs.columns.tolist()] + df_ncs.values.tolist())

    create_generic_table(document, rows_data, text, col_widths=[3, 3, 0.5, 0.3, 3, 3], cell_padding=0.2, align_left=False, font_size=8)

    
def format_header_cell(cell, text, font_size=10, font_name="Arial", bg_color="D9D9D9"):
    """Formata célula de cabeçalho: negrito, centralizado e fundo colorido."""
    para = cell.paragraphs[0]
    para.text = ""
    run = para.add_run(str(text))
    run.bold = True
    run.font.name = font_name
    run.font.size = Pt(font_size)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell._tc.get_or_add_tcPr().append(apply_background_color(bg_color))


def format_data_cell(cell, value, font_size=10, font_name="Arial"):
    """Formata célula de dados: centralizado e sem 'nan'."""
    para = cell.paragraphs[0]
    para.text = ""
    valor_texto = "" if pd.isna(value) else str(value)
    run = para.add_run(valor_texto)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER