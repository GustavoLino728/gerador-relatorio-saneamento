import os
import re
import sys
from PIL import Image, ImageOps
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Inches, Pt
from excel import get_non_conformities 
from utils import set_borders_table, get_images_from_dir, search_paragraph, sanitize_value
from paths import ASSETS_PATH, BASE_PATH


# Funções Utilitárias
def get_file_size_kb(file_path):
    """
    Retorna o tamanho de um arquivo em kb
    """
    return os.path.getsize(file_path) / 1024

def convert_to_valid_jpeg(image_path, target_size_kb=(20, 50), max_dimension=800):
        try:
            with Image.open(image_path) as img:
                if img.mode in ('RGBA', 'LA', 'P'):
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')

                img = ImageOps.exif_transpose(img)

                if max(img.size) > max_dimension:
                    img.thumbnail((max_dimension, max_dimension), Image.Resampling.LANCZOS)

                base_name = os.path.splitext(image_path)[0]
                new_path = f"{base_name}.jpg"

                if image_path != new_path and os.path.exists(image_path):
                    os.remove(image_path)

                quality = 85
                min_quality = 30
                step = 5

                while quality >= min_quality:
                    img.save(new_path, 'JPEG', quality=quality, optimize=True)
                    size_kb = get_file_size_kb(new_path)

                    if target_size_kb[0] <= size_kb <= target_size_kb[1]:
                        # print(f"✅ {os.path.basename(new_path)} otimizada para {size_kb:.1f}KB (q={quality})")
                        return new_path

                    if size_kb > target_size_kb[1]:
                        quality -= step
                    else:
                        print(f"⚠️ {os.path.basename(new_path)} muito pequeno: {size_kb:.1f}KB (q={quality})")
                        return new_path

                img.save(new_path, 'JPEG', quality=min_quality, optimize=True)
                size_kb = get_file_size_kb(new_path)
                # print(f"❌ {os.path.basename(new_path)} salvo com qualidade mínima ({min_quality}), {size_kb:.1f}KB")
                return new_path

        except Exception as e:
            print(f"❌ Erro ao processar {image_path}: {e}")
            return None

#----------------------------------


def build_caption_map(df, col_img="Nome da Foto", col_unit="Unidade", col_desc="Não Conformidade"):
    """
    Lê cada linha do DataFrame e monta um dicionario de legendas:
    - Se há múltiplas fotos separadas por vírgula (ex: "Foto 01, Foto 02"), 
      associa todas à mesma legenda.
    - Ignora casos onde a célula está vazia, None ou seja "Sem Foto".
    """
    captions = {}
    for _, row in df.iterrows():
        image_field = row[col_img]

        if not isinstance(image_field, str) or sanitize_value(image_field) == "semfoto":
            continue
        image_names = [img.strip() for img in re.split(r'[,;]', image_field) if sanitize_value(img) != "semfoto" and img.strip() != ""]
        for image_name in image_names:
            captions[image_name] = f"{image_name} - {row[col_unit]}: {row[col_desc]}"
    return captions


def process_images(path=ASSETS_PATH):
    """
    Processa imagens para JPEGs com qualidade iterativa visando tamanho entre target_size_kb
    """

    total = 0
    success = 0
    for root, dirs, files in os.walk(path):
        for filename in files:
            if filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                image_path = os.path.join(root, filename)
                total += 1
                if convert_to_valid_jpeg(image_path):
                    success += 1

    # print(f"Processamento finalizado: {success}/{total} imagens otimizadas.")


def validate_image(img_path):
    """
    Valida se a imagem pode ser aberta pelo PIL e python-docx
    """
    try:
        with Image.open(img_path) as img:
            img.verify()

        if os.path.getsize(img_path) == 0:
            return False, "Arquivo vazio"

        return True, "OK"
    except Exception as e:
        return False, str(e)


def create_table_images(document, insert_coord, list_of_images_path, captions=None, title_text="Registros Fotográficos"):
    """
    Cria tabela de imagens com legenda.
    """
    valid_images = []

    for img_path in list_of_images_path:

        abs_img_path = os.path.join(BASE_PATH, img_path)
        is_valid, error_msg = validate_image(abs_img_path)
        if is_valid:
            valid_images.append(abs_img_path)
        else:
            print(f"⚠️ Imagem inválida ignorada: {os.path.basename(img_path)} - {error_msg}")

    if not valid_images:
        print("❌ Nenhuma imagem válida encontrada para criar tabela")
        return insert_coord

    space_before = document.add_paragraph()
    space_after = document.add_paragraph()
    space_before_element = space_before._element
    space_after_element = space_after._element

    title = document.add_paragraph(title_text)
    title.style = 'Arial10'

    num_images = len(valid_images)
    num_rows = ((num_images + 1) // 2) * 2
    images_table = document.add_table(rows=num_rows, cols=2)

    for i, img_path in enumerate(valid_images):
        image_line = i // 2 * 2
        subtitle_line = image_line + 1
        column = i % 2

        try:
            images_table.cell(image_line, column).paragraphs[0].add_run().add_picture(
                img_path, width=Inches(3.3), height=Inches(2.5)
            )

            image_name = os.path.splitext(os.path.basename(img_path))[0]
            subtitle = captions.get(image_name, f"{image_name} - SEM LEGENDA") if captions else image_name

            cell = images_table.cell(subtitle_line, column)
            paragraph = cell.paragraphs[0]
            paragraph.style = 'Arial10'
            run = paragraph.add_run(subtitle)
            run.font.name = 'Arial'
            run.font.size = Pt(10)

        except UnrecognizedImageError:
            print(f"❌ Erro ao inserir imagem: {os.path.basename(img_path)} - Formato não reconhecido")
            error_cell = images_table.cell(image_line, column)
            error_paragraph = error_cell.paragraphs[0]
            error_run = error_paragraph.add_run(f"ERRO: {os.path.basename(img_path)}")
            error_run.font.name = 'Arial'
            error_run.font.size = Pt(10)

        except Exception as e:
            print(f"❌ Erro inesperado ao inserir {os.path.basename(img_path)}: {str(e)}")

    insert_coord._element.addnext(space_before_element)
    space_before_element.addnext(title._element)
    title._element.addnext(space_after_element)
    space_after_element.addnext(images_table._element)

    set_borders_table(images_table)
    return images_table


def divide_images(document, insert_coord, list_of_images_path, captions=None, block_size=6):
    """
    Divide imagens em blocos de N e cria tabelas
    """
    for i in range(0, len(list_of_images_path), block_size):
        table_images = list_of_images_path[i:i + block_size]
        insert_coord = create_table_images(document, insert_coord, table_images, captions)
    return insert_coord


def create_all_appendix_images(document, text_nc):
    """
    Cria todas tabelas de imagens para os apêndices, processando as imagens primeiro
    """
    process_images(ASSETS_PATH)

    images_by_folder = get_images_from_dir(ASSETS_PATH)

    if "fotos_nao_conformidades" in images_by_folder:
        df_ncs = get_non_conformities()
        captions_nc = build_caption_map(df_ncs)
        divide_images(document, text_nc, images_by_folder["fotos_nao_conformidades"], captions=captions_nc, block_size=6)

    if "fotos_condicoes_gerais" in images_by_folder:
        text_info = document.paragraphs[search_paragraph(document, "APÊNDICE 2 – CONDIÇÕES GERAIS")[-1]]
        divide_images(document, text_info, images_by_folder["fotos_condicoes_gerais"], captions=None, block_size=6)