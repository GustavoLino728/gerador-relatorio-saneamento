from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from unidecode import unidecode
from excel import get_non_conformities, get_inspections_data, list_of_all_units, documents_excel, town_statistics
from utils import search_paragraph, apply_background_color, set_column_widths, format_dict_values, set_table_margins, sanitize_value, set_borders_table, to_rows_data
import pandas as pd


def create_generic_table(document, rows_data, text_after_paragraph, col_widths=None,
                         cell_padding=0.1, align_left=False, font_size=10):
    """
    Cria uma tabela genérica a partir de rows_data (lista de listas).

    rows_data:
        - Primeira linha: cabeçalho (colunas)
        - Demais linhas: dados
        - Se uma linha tiver menos valores que o número de colunas ou apenas 1 valor:
          será tratada como subtítulo (mescla todas as colunas)
    text_after_paragraph: posição para inserir a tabela
    col_widths: lista de larguras para cada coluna (ex: [2, 6, 3])
    cell_padding: preenchimento interno das células
    align_left: se True, alinha conteúdo à esquerda
    font_size: tamanho da fonte do conteúdo
    """
    if not rows_data:
        print("⚠️ Nenhum dado para criar tabela.")
        return

    n_cols = max(len(row) for row in rows_data)
    table = document.add_table(rows=0, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows_data):
        row_cells = table.add_row().cells

        # Linha subtítulo / merge
        if len(row) == 1 or len(row) < n_cols:
            merged_cell = row_cells[0]
            for c in row_cells[1:]:
                merged_cell = merged_cell.merge(c)
            format_header_cell(merged_cell, row[0], font_size=font_size)
            set_table_margins(merged_cell, top=cell_padding, bottom=cell_padding,
                              start=cell_padding, end=cell_padding)

        elif i == 0:
            for j, value in enumerate(row):
                format_header_cell(row_cells[j], value, font_size=font_size)
                set_table_margins(row_cells[j], top=cell_padding, bottom=cell_padding,
                                  start=cell_padding, end=cell_padding)

        else:
            for j, value in enumerate(row):
                format_data_cell(row_cells[j], value, font_size=font_size)
                set_table_margins(row_cells[j], top=cell_padding, bottom=cell_padding,
                                  start=cell_padding, end=cell_padding)

    if align_left:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    set_borders_table(table)
    if col_widths:
        set_column_widths(table, *col_widths)

    paragraph_index = search_paragraph(document, text_after_paragraph)[0]
    document.paragraphs[paragraph_index]._element.addnext(table._element)


def create_abbreviations_table(document, text):
    """
    Cria a tabela de siglas mostrando o significado de cada sigla que pode aparecer no relatorio (FIXO)
    Formato: 20 linhas x 2 colunas.
    Linha 1: Cabeçalho (Sigla, Definição)
    Linha 2-20: Siglas e Seus Significados
    """
    abbreviations = [
        ["Sigla", "Definição"],
        ["ETA", "Estação de Tratamento de Água"],
        ["ETE", "Estação de Tratamento de Esgoto"],
        ["EEab", "Estação Elevatória de água bruta"],
        ["EEat", "Estação Elevatória de água tratada"],
        ["REL", "Reservatório Elevado"],
        ["RAP", "Reservatório Apoiado"],
        ["CMB", "Conjunto Moto Bomba"],
        ["GNR", "Gerência de Unidade de Negócios Regional"],
        ["SAA", "Sistemas de Abastecimento de Água"],
        ["SES", "Sistemas de Esgotamento Sanitário"],
        ["IUA", "Índice de Universalização do Abastecimento de Água"],
        ["IUE", "Índice de Universalização de Coleta de Esgotos Sanitários"],
        ["IUT", "Índice de Universalização de Tratamento de Esgotos Sanitários"],
        ["ICA", "Índice de Cobertura de Abastecimento de Água"],
        ["ICE", "Índice de Cobertura de Esgotamento Sanitário"],
        ["IPD", "Índice de Perdas na Distribuição"],
        ["IQAP", "Índice da Qualidade da Água Potável"],
        ["NBR", "Normas Brasileiras"],
        ["CSAN", "Coordenadoria de Saneamento da ARPE"]
    ]
    
    create_generic_table(document, abbreviations, text, col_widths=[1, 9], align_left=True)
    
    
def create_general_information_table(document, text):
    """
    Cria a tabela de Informações gerais da fiscalização. Sobre o regulador, o regulado e o titular
    Formato: 17 linhas x 2 colunas.
    Linha 1/6/11: Titulos (3.1 DO TITULAR, 3.2 DO REGULADO, 3.3 DO REGULADOR)
    """
    report_data = get_inspections_data()
    
    general_info = [
        ["3.1 DO TITULAR"],
        ["Titular:", "Microrregião de Água e Esgoto RMR-PAJEÚ/Microrregião de Água e Esgoto SERTÃO"],
        ["Endereço:", "Avenida Cruz Cabugá, 1387 - Santo Amaro - Recife, PE - CEP: 50040-905"],
        ["Responsável:", "Artur Paiva Coutinho"],
        ["Município:", report_data["Municipio"]],
        ["3.2 DO REGULADO"],
        ["Regulado:", "Companhia Pernambucana de Saneamento - Compesa"],
        ["Responsável:", "Dr. Alex Machado Campos"],
        ["Endereço:", "Av. Cruz Cabugá, 1387 - Santo Amaro - Recife, PE - CEP: 50040-905"],
        ["Representantes por acompanhar:", report_data["Representantes por acompanhar"]],
        ["3.3 DO REGULADOR"],
        ["Regulador:", "Agência de Regulação de Pernambuco"],
        ["Diretor Presidente:", "Carlos Porto Filho"],
        ["Endereço:", "Avenida Conselheiro Rosa e Silva, 975, Aflitos, Recife/PE, CEP: 52.050-020."],
        ["Responsáveis pela fiscalização:", f"{report_data['Analista 1']} e {report_data['Analista 2']}"],
        ["Período da Fiscalização:", report_data["Periodo"]],
        ["Tipo de Fiscalização:", "Direta e periódica."]
    ]
    
    create_generic_table(document, rows_data=general_info, text_after_paragraph=text, col_widths=[1, 9], align_left=True)


def create_documents_table(document, text):
    """
    Cria a tabela 1 - relativa as documentações necessarias para o processo, se elas foram enviadas ou não e porquê
    Formato: 12 linhas x 4 colunas.
    Linha 1: cabeçalho (DOCUMENTAÇÃO, SIM, NÃO, OBSERVAÇÕES)
    Linha 2-12: valores correspondentes
    """
    df_documents = documents_excel.copy()
    table = document.add_table(rows=1, cols=len(df_documents.columns))
    set_column_widths(table, 6.5, 0.5, 0.5, 6.5)

    for idx, col_name in enumerate(df_documents.columns):
        format_header_cell(table.rows[0].cells[idx], col_name, font_size=10)
 
    for _, row in df_documents.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            format_data_cell(cells[idx], value, font_size=10)

    set_borders_table(table)
    paragraph_index = search_paragraph(document, text)[0]
    document.paragraphs[paragraph_index]._element.addnext(table._element)


def create_town_units_table(document, text):
    """
    Cria a tabela 2 - Lista de Todas as Unidades do Municipio, de acordo com o tipo da fiscalização, puxando da planilha (Lista-SES-e-SAA) que vem da compesa
    Formato: 1+N linhas x 4 colunas.
    Linha 1: cabeçalho (ITEM, SISTEMA, UNIDADE, OBSERVAÇÃO)
    Linha 2-8: valores correspondentes
    """
    report_data = get_inspections_data()
    town_name = sanitize_value(report_data["Municipio"])
    inspection_type = sanitize_value(report_data["Tipo da Fiscalização"])

    units_df = list_of_all_units.copy()
    units_df.columns = units_df.columns.str.strip()
    units_df["TOWN_NORMALIZED"] = units_df["MUNICÍPIO"].apply(sanitize_value)
    units_df["WATER_SEWER_NORMALIZED"] = units_df["ÁGUA/ESGOTO"].apply(sanitize_value)

    filtered_units = units_df[units_df["TOWN_NORMALIZED"] == town_name]

    if "esgoto" in inspection_type.lower():
        filtered_units = filtered_units[filtered_units["WATER_SEWER_NORMALIZED"].str.contains("esgoto")]
    else:
        filtered_units = filtered_units[filtered_units["WATER_SEWER_NORMALIZED"].str.contains("agua")]

    if filtered_units.empty:
        print("⚠️ Nenhuma unidade encontrada para este município/tipo de fiscalização.")
        return

    filtered_units.insert(0, "ITEM", range(1, len(filtered_units) + 1))
    filtered_units["OBSERVAÇÃO"] = ""

    final_df = filtered_units[["ITEM", "SISTEMA", "UNIDADE", "OBSERVAÇÃO"]]

    rows_data = [final_df.columns.tolist()]  
    for row in final_df.itertuples(index=False, name=None):
        rows_data.append(list(row))

    create_generic_table(document=document, rows_data=rows_data, text_after_paragraph=text, col_widths=[0.3, 4, 6, 1.7], align_left=True)


def create_last_report_table(document, text):
    """
    Cria a tabela 3 - Com informações relevantes a última fiscalização naquele municipio
    Formato: 5 linhas x 2 colunas.
    Linha 1: cabeçalho (Contexto)
    Linha 2: Ultima Fiscalização
    Linha 3: Total NCs da Ultima Fiscalização
    Linha 4: DESDOBRAMENTOS
    Linha 5: NCs RESIDUAIS
    """
    report_data = format_dict_values(get_inspections_data())
    last_report_data = {
        "ÚLTIMA FISCALIZAÇÃO": report_data["Ultima Fiscalização"],
        "TOTAL DE NCs DA ÚLTIMA FISCALIZAÇÂO": report_data["Total NCS UF"],
        "DESDOBRAMENTOS": report_data["Desdobramentos"],
        "NCs RESIDUAIS": report_data["NCS Residuais"]
    }
    
    rows = to_rows_data(last_report_data, subtitle="CONTEXTO")
    create_generic_table(document, rows, text, col_widths=[2.5, 5], cell_padding=0.4, align_left=True)


def create_statistics_table(document, text):
    """
    Cria a tabela 4 - de Informações sobre Pernambuco em Geral (Fixa) e o Municipio da fiscalização (que é retirado da planilha - Estatisticas). 
    Puxa os dados sobre EIA, EAE, EIE, EAT, EIT, DAP
    Formato: 8 linhas x 3 colunas.
    Linha 1: cabeçalho (INFORMAÇÃO, PERNAMBUCO, "Município")
    Linha 2-8: valores correspondentes
    """
    df_statistics = town_statistics.copy()
    report_data = get_inspections_data()
    town_name = report_data["Municipio"].upper()

    pernambuco_stats = {
        "Quantidade de economias residenciais ativas de água (A) - EAA": "2.261.695",
        "Quantidade de economias residenciais inativas de água (B)-EIA": "377.745",
        "Quantidade de economias residenciais ativas de esgoto (C) - EAE": "654.143",
        "Quantidade de economias residenciais inativas de esgoto (D) - EIE": "300.683",
        "Quantidade de economias residenciais ativas com tratamento de esgoto (E) - EAT": "654.143",
        "Quantidade de economias residenciais inativas com tratamento de esgoto (F) - EIT": "300.683",
        "Quantidade de domicílios residenciais existentes na área de abrangência do prestador de serviços (G) - DAP": "2.646.895"
    }

    columns = list(pernambuco_stats.keys())
    town_stats = df_statistics[df_statistics["ANO BASE: 2023"] == town_name]
    town_stats = town_stats[columns]

    rows_data = [["INFORMAÇÃO", "PERNAMBUCO", town_name]]
    for column in columns:
        town_value = ""
        if not town_stats.empty:
            town_value = town_stats.iloc[0][column]
            if isinstance(town_value, (int, float)):
                town_value = str(int(town_value))
        rows_data.append([column, pernambuco_stats[column], town_value])

    create_generic_table(document=document, rows_data=rows_data, text_after_paragraph=text, col_widths=[6, 1.5, 1.5], align_left=True, font_size=10)


def create_quality_index_table(document, text):
    """
    Cria a tabela 5 - de indicadores de qualidade para o município.
    Formato: 2 linhas x 8 colunas.
    Linha 1: cabeçalho ("Município" + indicadores)
    Linha 2: valores correspondentes
    """
    df_statistics = town_statistics.copy()
    report_data = get_inspections_data()
    report_town = report_data["Municipio"].upper()

    df_statistics.columns = df_statistics.columns.str.replace(r"\s+\(\%\)", "(%)", regex=True)
    df_statistics.columns = df_statistics.columns.str.strip()

    selected_columns = ["IUA(%)", "IUE(%)", "IUT(%)", "ICA", "ICE", "IPD", "IQAP"]

    stats_from_this_town = df_statistics[df_statistics["ANO BASE: 2023"] == report_town]

    valores_tabela = []
    for col in selected_columns:
        valor = "-"
        if not stats_from_this_town.empty and col in stats_from_this_town.columns:
            valor = stats_from_this_town.iloc[0][col]
            if isinstance(valor, (int, float)):
                valor = str(int(valor))
        valores_tabela.append(valor)

    rows_data = [["Município"] + selected_columns, [report_town] + valores_tabela]

    create_generic_table(document, rows_data, text, col_widths=[3] + [1]*7, align_left=True, font_size=10)


def create_non_conformities_table(document, text):
    """
    Cria a tabela de não conformidades.
    Formato: 1+N linhas x 6 colunas.
    Linha 1: cabeçalho (Unidade, Não Conformidad, Nome da Foto, Artigo, Enquadramento, Determinações)
    Linha 2-N: Dados correspondentes
    """
    df_ncs = get_non_conformities()

    selected_columns = [
        "Unidade", 
        "Não Conformidade", 
        "Nome da Foto", 
        "Artigo", 
        "Enquadramento", 
        "Determinações"
    ]

    df_ncs = df_ncs[selected_columns].fillna("")
    rows_data = to_rows_data([df_ncs.columns.tolist()] + df_ncs.values.tolist())

    create_generic_table(document, rows_data, text, col_widths=[3, 3, 0.5, 0.3, 3, 3], cell_padding=0.2, align_left=False, font_size=8)


def create_water_params_table(document, text):
    """
    Cria a tabela de Parametros de Agua, referente a tabela 7 do relatorio de agua. Aparece caso haja pelo menos uma unidade ETA entre as NCs. 
    Lista essas unidades e o analista preenche os campos (CLORO (mg.L ), TURBIDEZ (NTU), OBSERVAÇÕES)
    Formato: 1+N linhas x 4 colunas.
    Linha 1: cabeçalho (QUALIDADE DA ÁGUA (UNIDADES), CLORO (mg.L ), TURBIDEZ (NTU), OBSERVAÇÕES)
    Linha 2-N: Dados correspondentes
    """
    df_ncs = get_non_conformities()
    df_eta = df_ncs[df_ncs["Sigla"] == "ETA"]

    if df_eta.empty:
        return
    
    columns = ["QUALIDADE DA ÁGUA (UNIDADES)", "CLORO (mg.L )", "TURBIDEZ (NTU)", "OBSERVAÇÕES"]
    rows_data = [columns]

    for _, row in df_eta.iterrows():
        row_list = [
            row["Unidade"],  
            "",              
            "",             
            ""               
        ]
        rows_data.append(row_list)

    create_generic_table(document=document, rows_data=rows_data, text_after_paragraph=text, col_widths=[5, 1.5, 1.5, 3], align_left=False)
    
def create_sewage_params_table(document, text):
    """
    Cria a tabela de Parametros de Qualidade do Efluente. Referente a tabela 7 do relatorio de esgoto. Aparece caso haja pelo menos uma unidade ETE entre as NCs. Lista essas unidades e o analista preenche os campos 
    Formato: 1+N linhas x 3 colunas.
    Linha 1: cabeçalho (QUALIDADE DO EFLUENTE (UNIDADES), DBO filtrada(mg O2/L), OBSERVAÇÕES)
    Linha 2-N: Dados correspondentes
    """
    df_ncs = get_non_conformities()
    df_ete = df_ncs[df_ncs["Sigla"] == "ETE"]

    if df_ete.empty:
        return
    
    columns = ["QUALIDADE DO EFLUENTE (UNIDADES)", "DBO filtrada (mg O2/L)", "OBSERVAÇÕES"]
    rows_data = [columns]

    for _, row in df_ete.iterrows():
        row_list = [
            row["Unidade"],  
            "",              
            ""               
        ]
        rows_data.append(row_list)

    create_generic_table(document=document, rows_data=rows_data, text_after_paragraph=text, col_widths=[5, 2, 3], align_left=False)

def create_table_7(document):
    """Decide qual das tabelas 7 deve ser gerada com base no tipo da fiscalização"""
    report_data = get_inspections_data()
    inspection_type = sanitize_value(report_data["Tipo da Fiscalização"])
    if inspection_type == 'agua':
        create_water_params_table(document, "Tabela 7 - Parâmetros da qualidade da água.")
    elif inspection_type == 'esgoto':
        create_water_params_table(document, "Tabela 7 - Parâmetro(s) da qualidade do efluente.")
    else:
        print("❌ Tipo de Fiscalização não válido, insira um válido: Agua ou Esgoto")
    
def format_header_cell(cell, text, font_size=10, font_name="Arial", bg_color="D9D9D9"):
    """Formata célula de cabeçalho: negrito, centralizado e fundo colorido."""
    para = cell.paragraphs[0]
    para.text = ""
    run = para.add_run(str(text))
    run.bold = True
    run.font.name = font_name
    run.font.size = Pt(font_size)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell._tc.get_or_add_tcPr().append(apply_background_color(bg_color))


def format_data_cell(cell, value, font_size=10, font_name="Arial"):
    """Formata célula de dados: centralizado e sem 'nan'."""
    para = cell.paragraphs[0]
    para.text = ""
    valor_texto = "" if pd.isna(value) else str(value)
    run = para.add_run(valor_texto)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER