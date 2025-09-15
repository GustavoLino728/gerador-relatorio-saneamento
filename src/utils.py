import os
from datetime import datetime, date
from excel import get_inspections_data
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from unidecode import unidecode
from paths import DATA_PATH, REPORTS_PATH, ASSETS_PATH


def next_filename():
    """Monta o nome do arquivo com o ID da fiscalização e caso houver arquivos já existentes incrementa em numero ao lado. EX: Relatório - ID 2 (1)"""

    inspections_data = get_inspections_data()
    id = str(int(inspections_data["ID da Fiscalização"])) 
    name = f"RELATÓRIO - ID {id}.docx"
    path = os.path.join(REPORTS_PATH, name)

    counter = 1
    while os.path.exists(path):
        name = f"RELATÓRIO - ID {id} ({counter}).docx"
        path = os.path.join(REPORTS_PATH, name)
        counter += 1
    
    return path


def get_images_from_dir(path=ASSETS_PATH):
    """
    Retorna as imagens organizadas por subpasta (apêndice).
    Estrutura de saída:
    {
        "fotos_nao_conformidades": ["./assets/fotos_nao_conformidades/nc1.jpg", "./assets/fotos_nao_conformidades/nc2.png"],
        "fotos_condicoes_gerais": ["./assets/fotos_condicoes_gerais/geral1.jpg", "./assets/fotos_condicoes_gerais/geral2.png"]
    }
    """
    result = {}
    for root, dirs, files in os.walk(path):
        folder_name = os.path.basename(root)
        images = [os.path.join(root, f) for f in files if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
        if images:
            result[folder_name] = images
    return result


def search_paragraph(document, text):
    """
    Busca no documento word o texto passado como parametro e retorna a posição dele no documento
    document: arquivo (objeto)
    text: texto desejado para buscar
    """
    paragraphs_found_by_search = []
    for i, p in enumerate(document.paragraphs):
        if text in p.text:
            paragraphs_found_by_search.append(i)
    return paragraphs_found_by_search

def format_value(value):
    """
    Formata um valor para a inserção nas tabelas:
    - Retira NaN e deixa em branco
    - Converte de Float pra Int
    - Retira hora da data e deixa no padrão (DD/MM/AA)
    """
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    if isinstance(value, (datetime, date)):
        return value.strftime("%d/%m/%Y")
    return str(value)

def format_dict_values(data: dict):
    """
    Formata um dict inteiro para inserção na tabela
    """
    return {k: format_value(v) for k, v in data.items()}

def sanitize_value(value):
    """
    Converte um valor em string limpa:
    - Remove espaços extras
    - Converte para minúsculo
    - Remove acentos
    """
    return unidecode(str(value).strip().lower())

# Funções Utilitarias para outras funções, não utilizar
def replace_in_paragraph(paragraph, replacements):
        full_text = paragraph.text
        replaced_any = False
        for placeholder, value in replacements.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, value)
                replaced_any = True

        if replaced_any and paragraph.runs:
            first_run = paragraph.runs[0]
            first_run.text = full_text
            first_run.font.color.rgb = RGBColor(0, 0, 0)
            for run in paragraph.runs[1:]:
                run.text = ''

def set_margin(tag, value, tcMar):
        if value is not None:
            margin = tcMar.find(qn(tag))
            if margin is None:
                margin = OxmlElement(tag)
                tcMar.append(margin)
            margin.set(qn('w:w'), str(int(value * 567))) 
            margin.set(qn('w:type'), 'dxa')

# ---------------------------------------------------------------

def substitute_placeholders(document):
    """
    Define um padrão de Strings no documento ({{x}}), e substitui no documento e nas tabelas, de acordo com o dicionário retornado com os dados da fiscalização.
    Caso houver chaves iguais as Strings realiza a troca. Ex: ({{nome}}) no documento, ele percorre o dicionario e caso aja a chave nome ele troca ({{nome}})
    pelo valor correspondente a chave nome.
    """
    excel_data = get_inspections_data()

    replacements = {f"{{{{{k}}}}}": format_value(v) for k, v in excel_data.items()}

    for p in document.paragraphs:
        replace_in_paragraph(p, replacements)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)

                
def apply_background_color(color_hex: str):
    """
    Cria um elemento de sombreado de célula com cor de fundo.
    Ex: 'D9D9D9' para cinza claro.
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    return shading_elm


def set_column_widths(table, *widths):
    """
    Ajusta a largura das colunas de uma tabela do python-docx.
    
    Parâmetros:
        table (docx.table.Table): Tabela a ser formatada.
        *widths (float): Larguras das colunas em polegadas. 
                         Passar um valor para cada coluna.
    """
    for row in table.rows:
        for col_idx, width in enumerate(widths):
            if col_idx < len(row.cells):
                row.cells[col_idx].width = Inches(width)
                
                
def set_table_margins(cell, top=None, start=None, bottom=None, end=None):
    """
    Define margens internas de cada célula de uma tabela (em Twips).
    top, start, bottom, end -> valores em cm ou None.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)

    set_margin('w:top', top, tcMar)
    set_margin('w:start', start, tcMar)
    set_margin('w:bottom', bottom, tcMar)
    set_margin('w:end', end, tcMar)


def set_borders_table(table):
    """
    Adiciona bordas visíveis a uma tabela do python-docx.
    Pode ser usada para qualquer tabela de documento Word.
    """
    tbl = table._element

    tblPr_list = tbl.xpath('./w:tblPr')
    if tblPr_list:
        tblPr = tblPr_list[0]
    else:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Cria as bordas
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')    # tipo da borda
        border.set(qn('w:sz'), '8')          # espessura
        border.set(qn('w:space'), '0')       # espaço entre borda e conteúdo
        border.set(qn('w:color'), '000000')  # cor preta
        tblBorders.append(border)

    tblPr.append(tblBorders)
    
    
def to_rows_data(data, subtitle=None):
    """
    Converte um dicionário ou lista de tuplas em lista de listas compatível com create_generic_table.
    
    Parâmetros:
    - data: dict ou lista de tuplas/listas [(chave, valor), ...]
    - subtitle: opcional, string que vira linha mesclada no topo
    
    Retorno:
    - rows_data: lista de listas pronta para create_generic_table
    """
    rows = []

    if subtitle:
        rows.append([subtitle])

    if isinstance(data, dict):
        for key, value in data.items():
            rows.append([key, value])

    elif isinstance(data, list):
        for item in data:
            rows.append(list(item))
    
    return rows


def insert_blank_lines(document, position_paragraph, n_lines=1):
    """
    Insere n linhas em branco após o parágrafo informado.
    """
    last_elem = position_paragraph
    for _ in range(n_lines):
        blank_paragraph = document.add_paragraph()
        blank_paragraph.add_run()
        last_elem._element.addnext(blank_paragraph._element)
        last_elem = blank_paragraph
    return last_elem


def insert_general_condition_section(document, text):
    """
    Insere tanto no sumário quanto o apêndice de fotos caso a pasta de informações gerais contenha alguma imagem
    document: Document
    text_summary: Texto para ter como refêrencia onde irá inserir as informações no sumário
    text_appendix: Texto para ter como refêrencia onde irá inserir as informações no apêndice
    """
    images_by_folder = get_images_from_dir()
    
    if "fotos_condicoes_gerais" in images_by_folder and images_by_folder["fotos_condicoes_gerais"]:
        pos_summary_idx, pos_appendix_idx = search_paragraph(document, text)
        position_summary = document.paragraphs[pos_summary_idx]
        position_appendix = document.paragraphs[pos_appendix_idx]

        summary_paragraph = document.add_paragraph()
        summary_run = summary_paragraph.add_run("APÊNDICE 2 – CONDIÇÕES GERAIS")
        summary_run.font.size = Pt(10)    

        position_summary._element.addnext(summary_paragraph._element)

        insert_position = insert_blank_lines(document, position_appendix, n_lines=3)

        appendix_paragraph = document.add_paragraph()
        appendix_run = appendix_paragraph.add_run("APÊNDICE 2 – CONDIÇÕES GERAIS")
        appendix_run.font.size = Pt(10) 
        appendix_run.bold = True         

        insert_position._element.addnext(appendix_paragraph._element)

def decide_report_type():
    report_data = get_inspections_data()
    if report_data is None:
        return None
    inspection_type = sanitize_value(report_data["Tipo da Fiscalização"])
    if inspection_type == "agua":
        return  Document(os.path.join(DATA_PATH, "RELATÓRIO_AGUA_MODELO.docx"))
    elif inspection_type == "esgoto":
        return Document(os.path.join(DATA_PATH, "RELATÓRIO_ESGOTO_MODELO.docx"))
    elif inspection_type == "comercial":
        return Document(os.path.join(DATA_PATH, "RELATÓRIO_COMERCIAL_MODELO.docx"))